from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "True_North_Engineering_Technical_Handover.docx"
COVER = ROOT / "assets" / "img" / "hero-industrial-doc.png"
Image.open(ROOT / "assets" / "img" / "hero-industrial.webp").convert("RGB").save(COVER)
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10.5)
for name, size, color in [("Title", 28, "0A1E33"), ("Heading 1", 17, "0A1E33"), ("Heading 2", 13, "C1440E")]:
    styles[name].font.name = "Calibri"
    styles[name].font.size = Pt(size)
    styles[name].font.color.rgb = RGBColor.from_string(color)

header = section.header.paragraphs[0]
header.text = "TRUE NORTH ENGINEERING | TECHNICAL HANDOVER"
header.style = styles["Normal"]
header.runs[0].font.size = Pt(8)
header.runs[0].font.color.rgb = RGBColor.from_string("5B6B7D")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("tnorthengineering.com | Client handover | 2026")
footer.runs[0].font.size = Pt(8)

doc.add_picture(str(COVER), width=Inches(6.8))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("TRUE NORTH ENGINEERING")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor.from_string("0A1E33")
p = doc.add_paragraph("Website, Analytics, Database and Deployment Guide")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(15)
p.runs[0].font.color.rgb = RGBColor.from_string("C1440E")
doc.add_paragraph("Prepared for client operations and future maintainers. Domain: tnorthengineering.com.")

def heading(text, level=1):
    doc.add_heading(text, level=level)

def bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

def steps(items):
    for item in items:
        doc.add_paragraph(item, style="List Number")

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("1C4066")

heading("1. What was delivered")
doc.add_paragraph("A responsive corporate engineering website, Supabase quote and analytics storage, a protected CRM dashboard, PWA support, SEO metadata, security headers, and a static deployment package.")
bullets([
    "Hero carousel, glass navigation, floating WhatsApp and assistant controls, contact form and responsive image-rich project work.",
    "First-party visitor, page-view, CTA and quote-start events with quote conversion reporting.",
    "Quote validation, honeypot spam protection, IP-based rate limits and anonymized IP hashes.",
    "Admin dashboard at /admin.html, protected by Supabase email/password authentication."
])

heading("2. Technology stack")
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
table.rows[0].cells[0].text = "Layer"
table.rows[0].cells[1].text = "Implementation"
for left, right in [
    ("Frontend", "Semantic HTML, modular CSS, ES modules, responsive layouts, WebP imagery and a service worker."),
    ("Application", "Static HTML/CSS/JavaScript modules served by Cloudflare Pages, Netlify or another static host."),
    ("Database", "Supabase PostgreSQL tables quote_requests and analytics_events with Row Level Security."),
    ("Analytics", "First-party Supabase events: page_view and tracked CTA interactions, shown in the admin overview."),
    ("Security", "CSP, HSTS, no-sniff, frame protection, referrer and permissions policies."),
    ("Deployment", "Static hosting with Supabase project configuration and no server process required.")
]:
    row = table.add_row().cells
    row[0].text, row[1].text = left, right

heading("3. Database and analytics")
doc.add_paragraph("Run supabase/schema.sql followed by supabase/migrations/002_admin_v2.sql in the Supabase SQL Editor. The project uses quote_requests for CRM leads and analytics_events for visitor and interaction activity.")
bullets([
    "Row Level Security allows public quote/event inserts while authenticated admin users can read and manage records.",
    "Quote requests retain name, email, project type, optional area, message, status and timestamp.",
    "Analytics are first-party and do not require Google Analytics, Clarity or a third-party analytics account. Those optional integrations remain disabled until IDs are added."
])
doc.add_paragraph("Use Supabase project backups and point-in-time recovery for database protection. Keep the anon key in the frontend only; never expose a service-role key.")

heading("4. Running the website")
doc.add_paragraph("For local preview, use the supplied static preview server. Keep Supabase configuration and admin access controlled in the Supabase dashboard.")
code("node local-server.mjs")
doc.add_paragraph("Open http://localhost:4173/ for the website. Configure the public Supabase URL and publishable key in assets/js/supabase-config.js.")

heading("5. Client admin dashboard")
doc.add_paragraph("The client opens https://tnorthengineering.com/admin.html after launch and signs in with the Supabase Auth user created for the operations team.")
steps([
    "Create or invite authorized users in Supabase Authentication > Users.",
    "Open /admin.html and sign in with the assigned email and password.",
    "Use Overview for quote totals, visitors, average area, project mix and activity.",
    "Use Quote CRM for search, filtering, CSV export, status changes, email copying and deletion."
])

heading("6. Deployment and domain checklist")
doc.add_paragraph("The project is ready for static hosting such as Cloudflare Pages, Netlify or Vercel with the Supabase project connected.")
steps([
    "Publish the project root with no build command.",
    "Run both Supabase SQL files before launch.",
    "Add the production domain to Supabase Authentication URL Configuration.",
    "In the DNS provider for tnorthengineering.com, point the root domain and www host to the chosen hosting provider using its exact A or CNAME records.",
    "Enable the host's HTTPS certificate, then set tnorthengineering.com as the canonical custom domain.",
    "Verify the public site, quote form, /admin.html, sitemap.xml, robots.txt and security headers."
])
doc.add_paragraph("Publishing the domain itself requires access to the domain registrar and the selected hosting account. Do not change DNS until the hosting provider supplies the correct records.")

heading("7. Content and operations handoff")
bullets([
    "Confirm the public phone number (+92 334 8767060), email, office address, WhatsApp number and social links before launch.",
    "Replace demonstration project photography and testimonials with approved client-owned material over time.",
    "Update the contact and LocalBusiness schema whenever the public contact details change.",
    "Run a database backup before site, server or database upgrades."
])

doc.save(OUT)
print(OUT)
